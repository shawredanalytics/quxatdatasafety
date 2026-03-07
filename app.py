import os
import sqlite3
from datetime import datetime
from io import BytesIO

import bcrypt
import pandas as pd
import plotly.express as px
import streamlit as st
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader

try:
    from docx import Document
    from docx.shared import Inches

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


try:
    from dotenv import load_dotenv
except ImportError:
    def load_dotenv():
        return None


load_dotenv()


# Use absolute path relative to script for DB
DB_PATH = os.getenv("DB_PATH", os.path.join(os.path.dirname(os.path.abspath(__file__)), "healthsecure.db"))
DEFAULT_ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "admin")
DEFAULT_ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "admin123")


ISO_CONTROLS = [
    "Access Control Policy",
    "User Access Provisioning",
    "Backup Schedule and Testing",
    "Offsite Backup Storage",
    "Database Encryption",
    "Endpoint Disk Encryption",
    "Incident Response Plan",
    "Incident Logging and Monitoring",
    "Vendor Risk Assessment",
    "Vendor Data Processing Agreements",
    "Asset Inventory of Medical Devices",
    "Asset Inventory of Servers and Workstations",
    "Business Continuity Plan",
    "Disaster Recovery Testing",
    "Secure Configuration Baselines",
]


ISO_CONTROL_DETAILS = {
    "Access Control Policy": {
        "clause": "ISO/IEC 27001:2022 Annex A 5.15",
        "description": "Formal documented access control policy covering all healthcare information systems and electronic health records.",
    },
    "User Access Provisioning": {
        "clause": "ISO/IEC 27001:2022 Annex A 5.18",
        "description": "Controlled process for granting, modifying and revoking user access rights based on job role and approval.",
    },
    "Backup Schedule and Testing": {
        "clause": "ISO/IEC 27001:2022 Annex A 8.13",
        "description": "Defined backup schedule and regular restoration tests for critical clinical and administrative data.",
    },
    "Offsite Backup Storage": {
        "clause": "ISO/IEC 27001:2022 Annex A 8.13",
        "description": "Secure storage of backup copies in an offsite location or logically isolated cloud environment.",
    },
    "Database Encryption": {
        "clause": "ISO/IEC 27001:2022 Annex A 8.24",
        "description": "Use of strong cryptographic controls to protect databases holding patient and clinical information.",
    },
    "Endpoint Disk Encryption": {
        "clause": "ISO/IEC 27001:2022 Annex A 8.24",
        "description": "Full disk or volume encryption for laptops, desktops and mobile devices used for healthcare operations.",
    },
    "Incident Response Plan": {
        "clause": "ISO/IEC 27001:2022 Annex A 5.26",
        "description": "Documented incident response procedures for detecting, responding to and learning from security incidents.",
    },
    "Incident Logging and Monitoring": {
        "clause": "ISO/IEC 27001:2022 Annex A 8.15",
        "description": "Centralised logging and monitoring of security events on critical healthcare applications and infrastructure.",
    },
    "Vendor Risk Assessment": {
        "clause": "ISO/IEC 27001:2022 Annex A 5.19",
        "description": "Assessment of information security risks related to third party and vendor services handling health data.",
    },
    "Vendor Data Processing Agreements": {
        "clause": "ISO/IEC 27001:2022 Annex A 5.20",
        "description": "Documented agreements with vendors defining security, privacy and data processing obligations.",
    },
    "Asset Inventory of Medical Devices": {
        "clause": "ISO/IEC 27001:2022 Annex A 5.9",
        "description": "Up to date inventory of network connected and standalone medical devices that process health information.",
    },
    "Asset Inventory of Servers and Workstations": {
        "clause": "ISO/IEC 27001:2022 Annex A 5.9",
        "description": "Comprehensive inventory of servers, workstations and virtual machines supporting healthcare operations.",
    },
    "Business Continuity Plan": {
        "clause": "ISO/IEC 27001:2022 Annex A 5.30",
        "description": "Business continuity plan covering continuity of clinical and administrative services during disruptions.",
    },
    "Disaster Recovery Testing": {
        "clause": "ISO/IEC 27001:2022 Annex A 5.30",
        "description": "Regular testing and review of disaster recovery capabilities for key healthcare systems.",
    },
    "Secure Configuration Baselines": {
        "clause": "ISO/IEC 27001:2022 Annex A 8.9",
        "description": "Defined secure configuration standards and baseline hardening for systems, applications and networks.",
    },
}


RANSOMWARE_QUESTIONS = [
    "Are backups offline?",
    "Is MFA enabled?",
    "Is EHR encrypted?",
    "Is antivirus deployed?",
    "Is network segmented?",
    "Is patching automated?",
    "Is phishing training conducted?",
    "Is admin access restricted?",
]


ISO_OPTIONS = {
    "Not Implemented (0)": 0,
    "Partially Implemented (50)": 50,
    "Implemented (100)": 100,
}


RANSOMWARE_OPTIONS = {
    "No (0)": 0,
    "Partial (50)": 50,
    "Yes (100)": 100,
}


RANSOMWARE_DETAILS = {
    "Are backups offline?": {
        "label": "Are backup copies stored offline or logically isolated from the main network?",
        "clause": "ISO/IEC 27001:2022 Annex A 8.13",
        "description": "Assesses whether backups are protected from ransomware by being offline, immutable or stored in a segregated environment.",
    },
    "Is MFA enabled?": {
        "label": "Is multi factor authentication enabled for remote access and privileged user accounts?",
        "clause": "ISO/IEC 27001:2022 Annex A 5.17",
        "description": "Checks whether more than one authentication factor is required for high risk and administrative access.",
    },
    "Is EHR encrypted?": {
        "label": "Is the electronic health record system and stored patient data encrypted at rest?",
        "clause": "ISO/IEC 27001:2022 Annex A 8.24",
        "description": "Evaluates whether cryptographic controls protect sensitive health records from unauthorised disclosure.",
    },
    "Is antivirus deployed?": {
        "label": "Is anti malware protection deployed and regularly updated on all relevant systems?",
        "clause": "ISO/IEC 27001:2022 Annex A 8.7",
        "description": "Covers deployment and maintenance of anti malware tools on servers, workstations and endpoints.",
    },
    "Is network segmented?": {
        "label": "Is the network segmented to separate critical healthcare systems from general user networks?",
        "clause": "ISO/IEC 27001:2022 Annex A 8.20",
        "description": "Looks at use of network segmentation and zoning to limit lateral movement of ransomware.",
    },
    "Is patching automated?": {
        "label": "Is there an automated or regularly executed process for applying security patches?",
        "clause": "ISO/IEC 27001:2022 Annex A 8.8",
        "description": "Checks whether operating systems and applications receive timely security updates.",
    },
    "Is phishing training conducted?": {
        "label": "Is regular phishing and cyber security awareness training conducted for all staff?",
        "clause": "ISO/IEC 27001:2022 Annex A 6.3",
        "description": "Assesses whether staff receive ongoing training to recognise and report phishing attempts.",
    },
    "Is admin access restricted?": {
        "label": "Is administrative access restricted to authorised personnel following least privilege principles?",
        "clause": "ISO/IEC 27001:2022 Annex A 5.18",
        "description": "Evaluates whether privileged accounts are tightly controlled, monitored and limited.",
    },
}


def get_connection():
    return sqlite3.connect(DB_PATH, check_same_thread=False)


def init_db():
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash BLOB NOT NULL
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS assessments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            control_name TEXT NOT NULL,
            score INTEGER NOT NULL,
            timestamp TEXT NOT NULL
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS ransomware (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question TEXT NOT NULL,
            score INTEGER NOT NULL,
            timestamp TEXT NOT NULL
        )
        """
    )

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS assessment_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            org_name TEXT NOT NULL,
            assessor_username TEXT NOT NULL,
            created_at TEXT NOT NULL
        )
        """
    )

    cursor.execute("PRAGMA table_info(assessments)")
    assessment_columns = [row[1] for row in cursor.fetchall()]
    if "session_id" not in assessment_columns:
        cursor.execute("ALTER TABLE assessments ADD COLUMN session_id INTEGER")

    cursor.execute("PRAGMA table_info(ransomware)")
    ransomware_columns = [row[1] for row in cursor.fetchall()]
    if "session_id" not in ransomware_columns:
        cursor.execute("ALTER TABLE ransomware ADD COLUMN session_id INTEGER")

    cursor.execute(
        "SELECT id FROM users WHERE username = ?",
        (DEFAULT_ADMIN_USERNAME,),
    )
    existing_admin = cursor.fetchone()

    if DEFAULT_ADMIN_PASSWORD:
        password_hash = bcrypt.hashpw(
            DEFAULT_ADMIN_PASSWORD.encode("utf-8"),
            bcrypt.gensalt(),
        )
        if not existing_admin:
            cursor.execute(
                "INSERT INTO users (username, password_hash) VALUES (?, ?)",
                (DEFAULT_ADMIN_USERNAME, password_hash),
            )
        else:
            cursor.execute(
                "UPDATE users SET password_hash = ? WHERE id = ?",
                (password_hash, existing_admin[0]),
            )

    conn.commit()
    conn.close()


def create_assessment_session(org_name: str, assessor_username: str) -> int:
    now = datetime.utcnow().isoformat()
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """
        INSERT INTO assessment_sessions (org_name, assessor_username, created_at)
        VALUES (?, ?, ?)
        """,
        (org_name.strip(), assessor_username.strip(), now),
    )
    session_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return session_id


def get_assessment_sessions() -> pd.DataFrame:
    conn = get_connection()
    df = pd.read_sql_query(
        """
        SELECT id, org_name, assessor_username, created_at
        FROM assessment_sessions
        ORDER BY created_at DESC, id DESC
        """,
        conn,
    )
    conn.close()
    return df


def get_session_by_id(session_id: int) -> dict | None:
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id, org_name, assessor_username, created_at
        FROM assessment_sessions
        WHERE id = ?
        """,
        (session_id,),
    )
    row = cursor.fetchone()
    conn.close()
    if not row:
        return None
    return {
        "id": row[0],
        "org_name": row[1],
        "assessor_username": row[2],
        "created_at": row[3],
    }


def delete_assessment_session(session_id: int):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM assessments WHERE session_id = ?", (session_id,))
    cursor.execute("DELETE FROM ransomware WHERE session_id = ?", (session_id,))
    cursor.execute("DELETE FROM assessment_sessions WHERE id = ?", (session_id,))
    conn.commit()
    conn.close()


def ensure_session_context() -> int | None:
    st.subheader("Assessment Session")
    sessions_df = get_assessment_sessions()
    current_session_id = st.session_state.get("current_session_id")

    selected_id = None
    if not sessions_df.empty:
        options = []
        id_map = {}
        for _, row in sessions_df.iterrows():
            label = f"{row['org_name']} - {row['created_at']}"
            options.append(label)
            id_map[label] = int(row["id"])

        default_index = 0
        if current_session_id is not None:
            for idx, label in enumerate(options):
                if id_map[label] == current_session_id:
                    default_index = idx
                    break

        selected_label = st.selectbox(
            "Select existing assessment",
            options,
            index=default_index,
            key="session_select",
        )
        selected_id = id_map[selected_label]
        st.session_state["current_session_id"] = selected_id
        session = get_session_by_id(selected_id)
        if session:
            st.session_state["current_org_name"] = session["org_name"]
    else:
        st.info("No assessments yet. Create a new one below.")

    new_org = st.text_input(
        "Register new organization for data safety assessment",
        key="new_org_name",
    )
    if st.button("Register organization and create assessment"):
        if not new_org.strip():
            st.error("Organization name is required for a new assessment.")
        else:
            assessor = st.session_state.get("username", "admin")
            new_id = create_assessment_session(new_org.strip(), assessor)
            st.session_state["current_session_id"] = new_id
            st.session_state["current_org_name"] = new_org.strip()
            st.success(f"Created new assessment for {new_org.strip()}.")
            selected_id = new_id

    if not sessions_df.empty:
        st.markdown("---")
        st.subheader("Delete registered organization and its assessments")
        delete_options = [
            f"{row['org_name']} - {row['created_at']}" for _, row in sessions_df.iterrows()
        ]
        delete_map = {
            f"{row['org_name']} - {row['created_at']}": int(row["id"])
            for _, row in sessions_df.iterrows()
        }
        delete_label = st.selectbox(
            "Select organization assessment to delete",
            delete_options,
            key="session_delete_select",
        )
        confirm_delete = st.checkbox(
            "I understand this will permanently delete this organization's assessment data.",
            key="session_delete_confirm",
        )
        if st.button("Delete selected organization assessment"):
            if not confirm_delete:
                st.warning("Please confirm deletion before proceeding.")
            else:
                delete_id = delete_map[delete_label]
                delete_assessment_session(delete_id)
                if st.session_state.get("current_session_id") == delete_id:
                    st.session_state["current_session_id"] = None
                    st.session_state["current_org_name"] = None
                st.success("Organization assessment deleted.")
                try:
                    st.rerun()
                except AttributeError:
                    st.experimental_rerun()

    return selected_id


def validate_login(username: str, password: str) -> bool:
    username = username.strip()
    if not username or not password:
        return False

    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT password_hash FROM users WHERE username = ?",
        (username,),
    )
    row = cursor.fetchone()
    conn.close()

    if not row:
        return False

    stored_hash = row[0]
    if isinstance(stored_hash, str):
        stored_hash = stored_hash.encode("utf-8")

    return bcrypt.checkpw(password.encode("utf-8"), stored_hash)


def save_iso_assessment(responses: dict, session_id: int):
    now = datetime.utcnow().isoformat()
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "DELETE FROM assessments WHERE session_id = ?",
        (session_id,),
    )
    for control, score in responses.items():
        cursor.execute(
            """
            INSERT INTO assessments (control_name, score, timestamp, session_id)
            VALUES (?, ?, ?, ?)
            """,
            (control, int(score), now, session_id),
        )
    conn.commit()
    conn.close()


def save_ransomware_assessment(responses: dict, session_id: int):
    now = datetime.utcnow().isoformat()
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "DELETE FROM ransomware WHERE session_id = ?",
        (session_id,),
    )
    for question, score in responses.items():
        cursor.execute(
            """
            INSERT INTO ransomware (question, score, timestamp, session_id)
            VALUES (?, ?, ?, ?)
            """,
            (question, int(score), now, session_id),
        )
    conn.commit()
    conn.close()


def load_iso_assessment_for_session(session_id: int):
    conn = get_connection()
    df = pd.read_sql_query(
        """
        SELECT control_name, score, timestamp
        FROM assessments
        WHERE session_id = ?
        ORDER BY id ASC
        """,
        conn,
        params=(session_id,),
    )
    conn.close()
    if df.empty:
        return None
    return df


def load_ransomware_assessment_for_session(session_id: int):
    conn = get_connection()
    df = pd.read_sql_query(
        """
        SELECT question, score, timestamp
        FROM ransomware
        WHERE session_id = ?
        ORDER BY id ASC
        """,
        conn,
        params=(session_id,),
    )
    conn.close()
    if df.empty:
        return None
    return df


def calculate_iso_score(df: pd.DataFrame | None) -> float:
    if df is None or df.empty:
        return 0.0
    return float(df["score"].mean())


def calculate_ransomware_score(df: pd.DataFrame | None) -> float:
    if df is None or df.empty:
        return 0.0
    return float(df["score"].mean())


def classify_ransomware_risk(score: float) -> tuple[str, str]:
    if score >= 80:
        return "Low Risk", "green"
    if 50 <= score < 80:
        return "Moderate Risk", "orange"
    return "High Risk", "red"


def build_iso_charts(df: pd.DataFrame | None):
    if df is None or df.empty:
        st.info("No ISO assessment data available yet.")
        return

    buckets = []
    for _, row in df.iterrows():
        if row["score"] >= 100:
            buckets.append("Implemented (100)")
        elif row["score"] >= 50:
            buckets.append("Partially Implemented (50)")
        else:
            buckets.append("Not Implemented (0)")

    df_levels = pd.DataFrame(
        {"control_name": df["control_name"], "implementation_level": buckets}
    )

    pie_data = (
        df_levels["implementation_level"]
        .value_counts()
        .reset_index(name="count")
        .rename(columns={"implementation_level": "level"})
    )

    pie_fig = px.pie(
        pie_data,
        names="level",
        values="count",
        title="ISO 27001:2022 Implementation Levels",
        color="level",
        color_discrete_map={
            "Implemented (100)": "green",
            "Partially Implemented (50)": "orange",
            "Not Implemented (0)": "red",
        },
    )
    st.plotly_chart(pie_fig, use_container_width=True)

    weak_df = df.sort_values("score").head(5)
    bar_fig = px.bar(
        weak_df,
        x="control_name",
        y="score",
        title="Weakest Controls",
        labels={"score": "Score"},
        color="score",
        color_continuous_scale=["red", "orange", "green"],
    )
    bar_fig.update_layout(xaxis_tickangle=-45)
    st.plotly_chart(bar_fig, use_container_width=True)


def generate_pdf_report(
    org_name: str,
    iso_score: float,
    ransomware_score: float,
    iso_df: pd.DataFrame | None,
    ransomware_df: pd.DataFrame | None,
    assessment_number: str,
):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    table_cell_style = ParagraphStyle(
        "TableCell",
        parent=styles["Normal"],
        fontSize=7,
        leading=8,
        wordWrap="CJK",
    )
    elements = []
    
    # Use relative path for logo
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png")
    
    if os.path.exists(logo_path):
        try:
            img = ImageReader(logo_path)
            iw, ih = img.getSize()
            aspect = ih / float(iw)
            width = 1.5 * 72
            height = width * aspect
            if height > 0.75 * 72: # Limit height
                height = 0.75 * 72
                width = height / aspect
            elements.append(Image(logo_path, width=width, height=height))
        except Exception:
            elements.append(Image(logo_path, width=1.5*72, height=0.5*72))
        elements.append(Spacer(1, 12))

    app_name = "QuXAT Data Safety Application"
    title = f"{app_name} - Assessment Report for {org_name}"
    elements.append(Paragraph(title, styles["Title"]))
    elements.append(Spacer(1, 12))

    date_str = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    elements.append(Paragraph(f"Assessment Date: {date_str}", styles["Normal"]))
    elements.append(Paragraph(f"Assessment Number: {assessment_number}", styles["Normal"]))
    elements.append(Spacer(1, 12))

    elements.append(
        Paragraph(f"ISO 27001:2022 Compliance Score: {iso_score:.1f}%", styles["Heading3"])
    )
    elements.append(
        Paragraph(
            f"Ransomware Risk Score: {ransomware_score:.1f}%",
            styles["Heading3"],
        )
    )
    elements.append(Spacer(1, 12))

    if iso_df is not None and not iso_df.empty:
        elements.append(Paragraph("ISO 27001:2022 Assessment Details", styles["Heading2"]))
        iso_rows = []
        for _, row in iso_df.iterrows():
            name = row["control_name"]
            details = ISO_CONTROL_DETAILS.get(name, {})
            clause = details.get("clause", "")
            score_value = int(row["score"])
            status_label = next(
                (label for label, value in ISO_OPTIONS.items() if value == score_value),
                str(score_value),
            )
            name_cell = Paragraph(str(name), table_cell_style)
            clause_cell = Paragraph(str(clause), table_cell_style)
            iso_rows.append([name_cell, clause_cell, score_value, status_label])

        iso_table_data = [["Control", "ISO 27001:2022 Clause", "Score", "Status"]] + iso_rows
        iso_table = Table(
            iso_table_data,
            hAlign="LEFT",
            colWidths=[
                doc.width * 0.34,
                doc.width * 0.27,
                doc.width * 0.09,
                doc.width * 0.20,
            ],
            repeatRows=1,
        )
        iso_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ALIGN", (2, 1), (2, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("FONTSIZE", (0, 0), (-1, 0), 7),
                    ("FONTSIZE", (0, 1), (-1, -1), 7),
                ]
            )
        )
        elements.append(iso_table)
        elements.append(Spacer(1, 12))

        not_impl = iso_df[iso_df["score"] == 0]
        not_impl_rows = [
            [row["control_name"], int(row["score"])] for _, row in not_impl.iterrows()
        ]
        if not_impl_rows:
            elements.append(Paragraph("Controls Not Implemented", styles["Heading2"]))
            table_data = [["Control", "Score"]] + not_impl_rows
            table = Table(
                table_data,
                hAlign="LEFT",
                colWidths=[260, 60],
            )
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ]
                )
            )
            elements.append(table)
            elements.append(Spacer(1, 12))

    if ransomware_df is not None and not ransomware_df.empty:
        elements.append(
            Paragraph("Ransomware Risk Assessment Details", styles["Heading2"])
        )
        r_rows = []
        for _, row in ransomware_df.iterrows():
            q_key = row["question"]
            details = RANSOMWARE_DETAILS.get(q_key, {})
            label = details.get("label", q_key)
            clause = details.get("clause", "")
            score_value = int(row["score"])
            answer_label = next(
                (label for label, value in RANSOMWARE_OPTIONS.items() if value == score_value),
                str(score_value),
            )
            label_cell = Paragraph(str(label), table_cell_style)
            clause_cell = Paragraph(str(clause), table_cell_style)
            r_rows.append([label_cell, clause_cell, score_value, answer_label])

        r_table_data = [["Question", "ISO 27001:2022 Clause", "Score", "Answer"]] + r_rows
        r_table = Table(
            r_table_data,
            hAlign="LEFT",
            colWidths=[
                doc.width * 0.42,
                doc.width * 0.20,
                doc.width * 0.09,
                doc.width * 0.19,
            ],
            repeatRows=1,
        )
        r_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ALIGN", (2, 1), (2, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("FONTSIZE", (0, 0), (-1, 0), 7),
                    ("FONTSIZE", (0, 1), (-1, -1), 7),
                ]
            )
        )
        elements.append(r_table)
        elements.append(Spacer(1, 12))

    elements.append(Paragraph("Recommendations", styles["Heading2"]))
    recommendations = [
        "Strengthen access control for critical healthcare systems and EHR.",
        "Ensure regular, tested offline backups for key clinical and administrative data.",
        "Implement network segmentation to limit ransomware spread.",
        "Roll out continuous phishing awareness training for staff.",
        "Review vendor security and data processing agreements regularly.",
    ]
    for rec in recommendations:
        elements.append(Paragraph(f"- {rec}", styles["Normal"]))

    doc.build(elements)
    pdf_value = buffer.getvalue()
    buffer.close()
    return pdf_value


def generate_word_report(
    org_name: str,
    iso_score: float,
    ransomware_score: float,
    iso_df: pd.DataFrame | None,
    ransomware_df: pd.DataFrame | None,
    assessment_number: str,
):
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not installed; Word report generation is unavailable.")

    buffer = BytesIO()
    document = Document()

    # Use relative path for logo
    current_dir = os.path.dirname(os.path.abspath(__file__))
    logo_filename = "logo.png"
    logo_path = os.path.join(current_dir, logo_filename)
    
    if os.path.exists(logo_path):
        document.add_picture(logo_path, width=Inches(1.5))

    app_name = "QuXAT Data Safety Application"
    document.add_heading(f"{app_name} - Assessment Report for {org_name}", level=1)

    date_str = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    document.add_paragraph(f"Assessment Date: {date_str}")
    document.add_paragraph(f"Assessment Number: {assessment_number}")
    document.add_paragraph("")

    document.add_heading("Summary Scores", level=2)
    document.add_paragraph(f"ISO 27001:2022 Compliance Score: {iso_score:.1f}%")
    document.add_paragraph(f"Ransomware Risk Score: {ransomware_score:.1f}%")
    document.add_paragraph("")

    if iso_df is not None and not iso_df.empty:
        document.add_heading("ISO 27001:2022 Assessment Details", level=2)
        iso_table = document.add_table(rows=1, cols=4)
        hdr_cells = iso_table.rows[0].cells
        hdr_cells[0].text = "Control"
        hdr_cells[1].text = "ISO 27001:2022 Clause"
        hdr_cells[2].text = "Score"
        hdr_cells[3].text = "Status"

        for _, row in iso_df.iterrows():
            name = row["control_name"]
            details = ISO_CONTROL_DETAILS.get(name, {})
            clause = details.get("clause", "")
            score_value = int(row["score"])
            status_label = next(
                (label for label, value in ISO_OPTIONS.items() if value == score_value),
                str(score_value),
            )
            cells = iso_table.add_row().cells
            cells[0].text = str(name)
            cells[1].text = str(clause)
            cells[2].text = str(score_value)
            cells[3].text = status_label

        document.add_paragraph("")

    if ransomware_df is not None and not ransomware_df.empty:
        document.add_heading("Ransomware Risk Assessment Details", level=2)
        r_table = document.add_table(rows=1, cols=4)
        hdr_cells = r_table.rows[0].cells
        hdr_cells[0].text = "Question"
        hdr_cells[1].text = "ISO 27001:2022 Clause"
        hdr_cells[2].text = "Score"
        hdr_cells[3].text = "Answer"

        for _, row in ransomware_df.iterrows():
            q_key = row["question"]
            details = RANSOMWARE_DETAILS.get(q_key, {})
            label = details.get("label", q_key)
            clause = details.get("clause", "")
            score_value = int(row["score"])
            answer_label = next(
                (label for label, value in RANSOMWARE_OPTIONS.items() if value == score_value),
                str(score_value),
            )
            cells = r_table.add_row().cells
            cells[0].text = str(label)
            cells[1].text = str(clause)
            cells[2].text = str(score_value)
            cells[3].text = answer_label

        document.add_paragraph("")

    document.add_heading("Recommendations", level=2)
    recommendations = [
        "Strengthen access control for critical healthcare systems and EHR.",
        "Ensure regular, tested offline backups for key clinical and administrative data.",
        "Implement network segmentation to limit ransomware spread.",
        "Roll out continuous phishing awareness training for staff.",
        "Review vendor security and data processing agreements regularly.",
    ]
    for rec in recommendations:
        document.add_paragraph(rec, style="List Bullet")

    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_blank_checklist_pdf():
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = styles["Title"]
    heading_style = styles["Heading2"]
    normal_style = styles["Normal"]
    
    elements = []
    
    # Use relative path for logo
    current_dir = os.path.dirname(os.path.abspath(__file__))
    logo_filename = "logo.png"
    logo_path = os.path.join(current_dir, logo_filename)
    
    if os.path.exists(logo_path):
        try:
            # Dynamically resize logo to fit within a reasonable box (e.g. 200x80)
            img = ImageReader(logo_path)
            iw, ih = img.getSize()
            aspect = ih / float(iw)
            width = 200
            height = width * aspect
            if height > 80:
                height = 80
                width = height / aspect
            elements.append(Image(logo_path, width=width, height=height))
        except Exception:
            elements.append(Image(logo_path, width=200, height=60))
        elements.append(Spacer(1, 12))
        
    # Title
    elements.append(Paragraph("QuXAT Healthcare Data Safety - Self Assessment Checklist", title_style))
    elements.append(Spacer(1, 12))
    
    elements.append(Paragraph("This checklist is designed to help organizations assess their data safety posture against ISO 27001:2022 controls and Ransomware readiness.", normal_style))
    elements.append(Spacer(1, 12))
    
    # Scoring Legend
    elements.append(Paragraph("Scoring Guide (Self-Assessment)", heading_style))
    elements.append(Paragraph("Use the following criteria to score your implementation status:", normal_style))
    elements.append(Spacer(1, 6))
    
    legend_data = [
        [
            Paragraph("<b>Score</b>", normal_style),
            Paragraph("<b>ISO 27001 Criteria</b>", normal_style),
            Paragraph("<b>Ransomware Readiness Criteria</b>", normal_style)
        ],
        [
            Paragraph("0", normal_style),
            Paragraph("Not Implemented - Control is missing or ad-hoc.", normal_style),
            Paragraph("No - Practice is not in place.", normal_style)
        ],
        [
            Paragraph("50", normal_style),
            Paragraph("Partially Implemented - Control exists but is not documented or consistently applied.", normal_style),
            Paragraph("Partial - Practice is in place but not fully effective or consistent.", normal_style)
        ],
        [
            Paragraph("100", normal_style),
            Paragraph("Implemented - Control is fully documented, implemented, and effective.", normal_style),
            Paragraph("Yes - Practice is fully established and effective.", normal_style)
        ]
    ]
    
    legend_table = Table(legend_data, colWidths=[40, 200, 200])
    legend_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(legend_table)
    elements.append(PageBreak())
    
    # ISO Section
    elements.append(Paragraph("ISO 27001:2022 Controls Assessment", heading_style))
    elements.append(Spacer(1, 12))
    
    iso_data = [[
        Paragraph("<b>Control</b>", normal_style),
        Paragraph("<b>Clause</b>", normal_style),
        Paragraph("<b>Description</b>", normal_style),
        Paragraph("<b>Score (0/50/100)</b>", normal_style)
    ]]
    for control in ISO_CONTROLS:
        details = ISO_CONTROL_DETAILS.get(control, {})
        clause = details.get("clause", "")
        desc = details.get("description", "")
        
        iso_data.append([
            Paragraph(control, normal_style),
            Paragraph(clause, normal_style),
            Paragraph(desc, normal_style),
            "" 
        ])
        
    iso_table = Table(iso_data, colWidths=[110, 70, 180, 80])
    iso_table.setStyle(TableStyle([
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(iso_table)
    elements.append(PageBreak())
    
    # Ransomware Section
    elements.append(Paragraph("Ransomware Readiness Checklist", heading_style))
    elements.append(Spacer(1, 12))
    
    ransom_data = [[
        Paragraph("<b>Question</b>", normal_style),
        Paragraph("<b>Clause</b>", normal_style),
        Paragraph("<b>Description</b>", normal_style),
        Paragraph("<b>Score (0/50/100)</b>", normal_style)
    ]]
    for question in RANSOMWARE_QUESTIONS:
        details = RANSOMWARE_DETAILS.get(question, {})
        label = details.get("label", question)
        clause = details.get("clause", "")
        desc = details.get("description", "")
        
        ransom_data.append([
            Paragraph(label, normal_style),
            Paragraph(clause, normal_style),
            Paragraph(desc, normal_style),
            ""
        ])
        
    ransom_table = Table(ransom_data, colWidths=[110, 70, 180, 80])
    ransom_table.setStyle(TableStyle([
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(ransom_table)
    
    elements.append(PageBreak())
    
    # Contact Details
    elements.append(Paragraph("<b>QuXAT Data Security - Advisory Team</b>", heading_style))
    elements.append(Spacer(1, 6))
    elements.append(Paragraph("<b>Whatsapp:</b> +91 6301237212", normal_style))
    elements.append(Paragraph("<b>Email ID:</b> quxat.team@gmail.com", normal_style))
    elements.append(Paragraph("<b>Website:</b> www.quxat.com", normal_style))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


def show_login():
    st.markdown(
        """
        <style>
        header[data-testid="stHeader"] {
            display: none;
        }
        .block-container {
            padding-top: 2rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Use relative path for logo
    current_dir = os.path.dirname(os.path.abspath(__file__))
    logo_filename = "logo.png"
    logo_path = os.path.join(current_dir, logo_filename)
    
    if os.path.exists(logo_path):
        st.image(logo_path, width=280)

    st.title("QuXAT Data Safety Application")
    st.subheader("Secure Healthcare Data with Ransomware Risk Dashboard")

    st.markdown("---")
    st.write("Interested in undergoing the QuXAT Data Safety Process? Download our checklist below:")
    
    pdf_bytes = generate_blank_checklist_pdf()
    st.download_button(
        label="Download Blank Assessment Checklist",
        data=pdf_bytes,
        file_name="QuXAT_Healthcare_Data_Safety_Self_Assessment_Checklist.pdf",
        mime="application/pdf"
    )
    st.markdown("---")

    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        if validate_login(username, password):
            st.session_state["authenticated"] = True
            st.session_state["username"] = username
            try:
                st.rerun()
            except AttributeError:
                st.experimental_rerun()
        else:
            st.error("Invalid username or password.")


def iso_assessment_page():
    st.header("ISO 27001:2022 Self-Assessment")
    st.write("Select the implementation status for each control.")
    session_id = ensure_session_context()
    if session_id is None:
        return

    session = get_session_by_id(session_id)
    if session:
        st.caption(f"Organization: {session['org_name']}  |  Assessor: {session['assessor_username']}")

    latest = load_iso_assessment_for_session(session_id)
    previous_scores = {}
    if latest is not None and not latest.empty:
        for _, row in latest.iterrows():
            previous_scores[row["control_name"]] = int(row["score"])

    responses = {}
    for control in ISO_CONTROLS:
        details = ISO_CONTROL_DETAILS.get(control, {})
        clause = details.get("clause", "")
        description = details.get("description", "")

        default_label = None
        if control in previous_scores:
            previous_score = previous_scores[control]
            for label, value in ISO_OPTIONS.items():
                if value == previous_score:
                    default_label = label
                    break

        label_text = f"{control} ({clause})" if clause else control
        selected = st.selectbox(
            label_text,
            list(ISO_OPTIONS.keys()),
            index=(
                list(ISO_OPTIONS.keys()).index(default_label)
                if default_label in ISO_OPTIONS
                else 0
            ),
            key=f"iso_{control}",
        )
        if description:
            st.caption(description)
        responses[control] = ISO_OPTIONS[selected]

    if st.button("Save ISO Assessment"):
        save_iso_assessment(responses, session_id)
        st.success("ISO 27001:2022 assessment saved.")


def ransomware_page():
    st.header("Ransomware Risk Checklist")
    st.write("Answer the questions to assess ransomware risk exposure.")
    session_id = ensure_session_context()
    if session_id is None:
        return

    session = get_session_by_id(session_id)
    if session:
        st.caption(f"Organization: {session['org_name']}  |  Assessor: {session['assessor_username']}")

    latest = load_ransomware_assessment_for_session(session_id)
    previous_scores = {}
    if latest is not None and not latest.empty:
        for _, row in latest.iterrows():
            previous_scores[row["question"]] = int(row["score"])

    responses = {}
    for question in RANSOMWARE_QUESTIONS:
        details = RANSOMWARE_DETAILS.get(question, {})
        label_text = details.get("label", question)
        clause = details.get("clause", "")
        description = details.get("description", "")

        default_label = None
        if question in previous_scores:
            previous_score = previous_scores[question]
            for option_label, value in RANSOMWARE_OPTIONS.items():
                if value == previous_score:
                    default_label = option_label
                    break

        full_label = f"{label_text} ({clause})" if clause else label_text
        selected = st.selectbox(
            full_label,
            list(RANSOMWARE_OPTIONS.keys()),
            index=(
                list(RANSOMWARE_OPTIONS.keys()).index(default_label)
                if default_label in RANSOMWARE_OPTIONS
                else 0
            ),
            key=f"ransom_{question}",
        )
        if description:
            st.caption(description)
        responses[question] = RANSOMWARE_OPTIONS[selected]

    if st.button("Save Ransomware Assessment"):
        save_ransomware_assessment(responses, session_id)
        st.success("Ransomware assessment saved.")


def dashboard_page():
    st.header("Dashboard")
    session_id = ensure_session_context()
    if session_id is None:
        return

    session = get_session_by_id(session_id)
    if session:
        st.caption(f"Organization: {session['org_name']}  |  Assessor: {session['assessor_username']}")

    iso_df = load_iso_assessment_for_session(session_id)
    ransomware_df = load_ransomware_assessment_for_session(session_id)

    iso_score = calculate_iso_score(iso_df)
    ransomware_score = calculate_ransomware_score(ransomware_df)
    risk_label, risk_color = classify_ransomware_risk(ransomware_score)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("ISO Compliance Score", f"{iso_score:.1f}%")
    with col2:
        st.metric("Ransomware Risk Score", f"{ransomware_score:.1f}%")
    with col3:
        st.markdown(
            f"<div style='padding:8px;border-radius:4px;background-color:{risk_color};color:white;text-align:center;'>"
            f"{risk_label}"
            "</div>",
            unsafe_allow_html=True,
        )

    st.subheader("ISO Implementation Overview")
    build_iso_charts(iso_df)


def assessments_history_page():
    st.header("Assessments History")
    sessions_df = get_assessment_sessions()
    if sessions_df.empty:
        st.info("No assessments have been created yet.")
        return

    summary_rows = []
    for _, row in sessions_df.iterrows():
        session_id = int(row["id"])
        iso_df = load_iso_assessment_for_session(session_id)
        ransomware_df = load_ransomware_assessment_for_session(session_id)
        iso_score = calculate_iso_score(iso_df)
        ransomware_score = calculate_ransomware_score(ransomware_df)
        assessment_number = f"HS-{session_id:05d}"
        summary_rows.append(
            {
                "Assessment Number": assessment_number,
                "Organization": row["org_name"],
                "Assessor": row["assessor_username"],
                "Created At": row["created_at"],
                "ISO Score": f"{iso_score:.1f}%",
                "Ransomware Score": f"{ransomware_score:.1f}%",
            }
        )

    summary_df = pd.DataFrame(summary_rows)
    st.dataframe(summary_df, use_container_width=True)

    labels = [
        f"HS-{int(r['id']):05d} - {r['org_name']}"
        for _, r in sessions_df.iterrows()
    ]
    label_to_id = {
        f"HS-{int(r['id']):05d} - {r['org_name']}": int(r["id"])
        for _, r in sessions_df.iterrows()
    }

    selected_label = st.selectbox(
        "Set current assessment",
        labels,
        key="history_session_select",
    )
    if st.button("Use this assessment"):
        selected_id = label_to_id[selected_label]
        st.session_state["current_session_id"] = selected_id
        session = get_session_by_id(selected_id)
        if session:
            st.session_state["current_org_name"] = session["org_name"]
        st.success("Current assessment updated. Navigate to Dashboard or other pages to view details.")

    st.markdown("---")
    st.subheader("Delete assessment")
    delete_label = st.selectbox(
        "Select assessment to delete",
        labels,
        key="delete_session_select",
    )
    confirm_delete = st.checkbox(
        "I understand this will permanently delete the selected assessment and all related data.",
        key="delete_confirm",
    )
    if st.button("Delete selected assessment"):
        if not confirm_delete:
            st.warning("Please confirm deletion before proceeding.")
        else:
            delete_id = label_to_id[delete_label]
            delete_assessment_session(delete_id)
            if st.session_state.get("current_session_id") == delete_id:
                st.session_state["current_session_id"] = None
                st.session_state["current_org_name"] = None
            st.success("Assessment deleted.")
            try:
                st.rerun()
            except AttributeError:
                st.experimental_rerun()


def report_page():
    st.header("Generate Report")
    session_id = ensure_session_context()
    if session_id is None:
        return

    session = get_session_by_id(session_id)
    org_name = session["org_name"] if session else "Unknown Organization"
    assessment_number = f"HS-{session_id:05d}"

    iso_df = load_iso_assessment_for_session(session_id)
    ransomware_df = load_ransomware_assessment_for_session(session_id)
    iso_score = calculate_iso_score(iso_df)
    ransomware_score = calculate_ransomware_score(ransomware_df)

    if iso_df is None or iso_df.empty:
        st.warning("Please complete and save an ISO assessment before generating a report.")
        return

    if ransomware_df is None or ransomware_df.empty:
        st.warning("Please complete and save a ransomware assessment before generating a report.")
        return

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Generate PDF"):
            pdf_bytes = generate_pdf_report(
                org_name,
                iso_score,
                ransomware_score,
                iso_df,
                ransomware_df,
                assessment_number,
            )
            st.download_button(
                label="Download PDF Report",
                data=pdf_bytes,
                file_name=f"healthsecure_report_{assessment_number}.pdf",
                mime="application/pdf",
            )
    with col2:
        if not HAS_DOCX:
            st.info("Word report download is unavailable because python-docx is not installed on the server.")
        else:
            if st.button("Generate Word"):
                docx_bytes = generate_word_report(
                    org_name,
                    iso_score,
                    ransomware_score,
                    iso_df,
                    ransomware_df,
                    assessment_number,
                )
                st.download_button(
                    label="Download Word Report",
                    data=docx_bytes,
                    file_name=f"healthsecure_report_{assessment_number}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )


def main():
    st.set_page_config(
        page_title="QuXAT Data Safety Application",
        layout="wide",
    )

    init_db()

    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False

    if not st.session_state["authenticated"]:
        show_login()
        return

    with st.sidebar:
        st.title("QuXAT Data Safety Application")
        st.write(f"Logged in as: {st.session_state.get('username', '')}")
        page = st.radio(
            "Navigation",
            ["Dashboard", "Assessments History", "ISO Assessment", "Ransomware Check", "Generate Report"],
        )
        if st.button("Logout"):
            st.session_state.clear()
            try:
                st.rerun()
            except AttributeError:
                st.experimental_rerun()

    if page == "Dashboard":
        dashboard_page()
    elif page == "Assessments History":
        assessments_history_page()
    elif page == "ISO Assessment":
        iso_assessment_page()
    elif page == "Ransomware Check":
        ransomware_page()
    elif page == "Generate Report":
        report_page()


if __name__ == "__main__":
    main()
